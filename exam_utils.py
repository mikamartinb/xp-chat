import io
import os
import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from model_utils import rag_generate_exam


def create_exam_document(JSON_DIR, vectorstore, page_name, output_DIR):
    """
    Create exam documents in DOCX (with and without answers) based on JSON configuration
    and a generated exam string via a RAG approach.

    This function performs the following steps:
      1. Load exam configuration data from a JSON file.
      2. Ensure that the output directories for DOCX and PDF files exist.
      3. Generate the exam text using a retrieval-augmented generation (RAG) method.
      4. Create two Word documents: one with correct answers and one without.
      5. Save the generated documents to the specified output directories.
    
    Args:
        JSON_DIR (str): Path to the JSON file containing exam configuration data.
        vectorstore: A vector store object used by the RAG exam generator.
        page_name (str): The base name for the exam documents.
        output_DIR (str): The directory where output folders (DOCX and PDF) will be created.
    
    Returns:
        bool: True if both documents are created and saved successfully, False otherwise.
    """
    # Load exam configuration data from JSON file
    with open(JSON_DIR) as json_file:
        form = json.load(json_file)

    # Ensure output directories for DOCX and PDF files exist
    docx_dir = os.path.join(output_DIR, "DOCX")
    pdf_dir = os.path.join(output_DIR, "PDF")
    if not os.path.exists(docx_dir):
        os.makedirs(docx_dir)
    if not os.path.exists(pdf_dir):
        os.makedirs(pdf_dir)

    # Define file names for exam documents (with and without answers)
    exam_with_docx = f"{page_name}_exam_with_answers.docx"
    exam_without_docx = f"{page_name}_exam_without_answers.docx"
    exam_with_DIR_docx = os.path.join(docx_dir, exam_with_docx)
    exam_without_DIR_docx = os.path.join(docx_dir, exam_without_docx)

    exam_with_pdf = f"{page_name}_exam_with_answers.pdf"
    exam_without_pdf = f"{page_name}_exam_without_answers.pdf"
    exam_with_DIR_pdf = os.path.join(pdf_dir, exam_with_pdf)
    exam_without_DIR_pdf = os.path.join(pdf_dir, exam_without_pdf)

    # Debug prints for file paths and vector store content
    print(f"Path for DOCX (with answers): {exam_with_DIR_docx}")
    print(f"Path for DOCX (without answers): {exam_without_DIR_docx}")
    print(f"Path for PDF (with answers): {exam_with_DIR_pdf}")
    print(f"Path for PDF (without answers): {exam_without_DIR_pdf}")
    print(f"Vector store content: {vectorstore}")

    # Print loaded form data for debugging
    print("Form data loaded from JSON:")
    for key, value in form.items():
        print(f"{key}: {value}")

    # Define logo path (adjust if necessary)
    logo_path = "public/leuphana_logo.png"

    # Generate exam string using RAG
    exam_string = rag_generate_exam(form, vectorstore)
    if exam_string:
        print("Exam generated successfully.")
    else:
        print("Error generating exam with JSON data or vectorstore")
        return False

    def create_document(filter_correct_answers=False):
        """
        Create a Word document containing the exam text.

        The document includes a header with a logo, title page, and formatted exam questions.
        If filter_correct_answers is True, the correct answers will be omitted from the document.

        Args:
            filter_correct_answers (bool): If True, filter out lines that display correct answers.
                                           Defaults to False.

        Returns:
            Document: A python-docx Document object with the formatted exam.
        """
        doc = Document()

        # Add logo to the header of the document
        section = doc.sections[0]
        header = section.header
        header_paragraph = header.paragraphs[0]
        header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # Center the logo
        run = header_paragraph.add_run()
        try:
            run.add_picture(logo_path, width=Inches(2))  # Insert logo with a width of 2 inches
        except Exception as e:
            print(f"Error when adding the logo to the header: {e}")

        # Create a title page by adding several empty paragraphs for spacing
        for _ in range(5):
            doc.add_paragraph()
        middle_paragraph = doc.add_paragraph()
        middle_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        middle_run = middle_paragraph.add_run(f"{form['module']}")
        middle_run.bold = True
        middle_run.font.size = Pt(24)
        for _ in range(7):
            doc.add_paragraph()

        # Add subtitle information (university, date, module, examiner, semester, and total points)
        subtitle_paragraph = doc.add_paragraph()
        subtitle_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        subtitle_text = (
            f"{form['university']}\n"
            f"{form['date']}\n"
            f"{form['module']}\n"
            f"{form['professor']}\n"
            f"{form['semester']}\n"
            f"Total: {form['num_tasks'] * form['num_points']} Pt."
        )
        subtitle_run = subtitle_paragraph.add_run(subtitle_text)
        subtitle_run.font.size = Pt(12)

        # Add a page break to separate the title page from the exam questions
        doc.add_page_break()

        # Set default document style
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)

        # Add an initial empty paragraph for spacing before questions
        empty_paragraphs_before_questions = 1
        for _ in range(empty_paragraphs_before_questions):
            doc.add_paragraph()

        # Process exam_string into lines and format them into the document
        lines = exam_string.split("\n")
        block = []  # Buffer for accumulating lines of a question or answer block
        question_count = 0

        for line in lines:
            line = line.strip()

            # If line indicates a question (wrapped in '**')
            if line.startswith("**") and line.endswith("**"):
                if block:
                    # Add the accumulated block as a paragraph
                    paragraph = doc.add_paragraph()
                    for b_line in block:
                        run = paragraph.add_run(b_line)
                        # Italicize if the line ends with a question mark
                        if b_line.endswith("?"):
                            run.italic = True
                        run.add_break()
                    paragraph.paragraph_format.keep_together = True
                    block = []
                    question_count += 1

                    # Insert a page break after every 4 questions
                    if question_count == 4:
                        doc.add_page_break()
                        question_count = 0
                    for _ in range(empty_paragraphs_before_questions):
                        doc.add_paragraph()

                # Add the question text in bold with points
                paragraph = doc.add_paragraph()
                question_text = line.strip("**")
                run = paragraph.add_run(f"{question_text} ({form['num_points']} Pt.)")
                run.bold = True
                paragraph.paragraph_format.keep_with_next = True

            # Process lines wrapped in '++' (could indicate section headers)
            elif line.startswith("++") and line.endswith("++"):
                if block:
                    paragraph = doc.add_paragraph()
                    for b_line in block:
                        run = paragraph.add_run(b_line)
                        if b_line.endswith("?"):
                            run.italic = True
                        run.add_break()
                    paragraph.paragraph_format.keep_together = True
                    block = []
                paragraph = doc.add_paragraph()
                text = line.strip("++")
                run = paragraph.add_run(text)
                run.bold = True

            # Accumulate lines that look like questions or answers
            elif line.endswith("?"):
                block.append(line)

            # Include correct answers only if not filtering them out
            elif line.startswith("Correct Answers:"):
                if not filter_correct_answers:
                    block.append(line)

            # Accumulate any non-empty lines
            elif line:
                block.append(line)

        # If there is any remaining block, add it to the document
        if block:
            paragraph = doc.add_paragraph()
            for b_line in block:
                run = paragraph.add_run(b_line)
                if b_line.endswith("?"):
                    run.italic = True
                run.add_break()
            paragraph.paragraph_format.keep_together = True

        return doc

    # Create document versions with and without answers
    doc_with_answers = create_document(filter_correct_answers=False)
    if doc_with_answers:
        print("Document with answers created successfully.")
    else:
        print("Failed to create document with answers.")
        return False

    doc_without_answers = create_document(filter_correct_answers=True)
    if doc_without_answers:
        print("Document without answers created successfully.")
    else:
        print("Failed to create document without answers.")
        return False

    # Save the generated DOCX files
    try:
        doc_with_answers.save(exam_with_DIR_docx)
        print(f"Document with answers saved successfully at {exam_with_DIR_docx}.")
    except Exception as e:
        print(f"Failed to save document with answers: {e}")
        return False

    try:
        doc_without_answers.save(exam_without_DIR_docx)
        print(f"Document without answers saved successfully at {exam_without_DIR_docx}.")
    except Exception as e:
        print(f"Failed to save document without answers: {e}")
        return False

    return True
